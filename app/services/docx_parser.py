"""
DOCX parser service for extracting meeting notes.
"""
from docx import Document
from typing import Dict, Optional, List
from datetime import datetime
import re


SECTION_HEADERS = [
    "General Discussion",
    "Gap Assessment",
    "Alignment Points",
    "Challenges",
    "Team",
    "Product / Pricing",
    "Customer",
    "Market / Competition",
    "P&L and Financials",
    "Gantt Update",
    "Key Action Items",
    "Feedback / Ideas"
]


def parse_single_meeting_docx(file_path: str) -> Dict:
    """
    Parse a DOCX file containing ONE meeting's notes.
    
    In production, users upload one meeting per file.
    For testing with Meeting_Dump.docx, call parse_multi_meeting_docx instead.
    
    Returns:
        Dict with {date, raw_text, sections}
    """
    doc = Document(file_path)
    raw_text = ""
    sections = {}
    current_section = None
    meeting_date = None
    
    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        
        raw_text += text + '\n'
        
        # Extract date if found (e.g., "06 Mar 26")
        if not meeting_date:
            date_match = re.match(r'^\d{2}\s\w{3}\s\d{2}$', text)
            if date_match:
                meeting_date = parse_date(text)
                continue
        
        # Check if it's a section header
        if text in SECTION_HEADERS:
            current_section = text
            sections[current_section] = []
            continue
        
        # Add content to current section
        if current_section:
            sections[current_section].append(text)
    
    return {
        'date': meeting_date,
        'raw_text': raw_text,
        'sections': sections
    }


def parse_multi_meeting_docx(file_path: str) -> List[Dict]:
    """
    Parse Meeting_Dump.docx containing MULTIPLE meetings (for testing).
    
    This is used ONLY for testing with Meeting_Dump.docx which has 9 meetings.
    In production, use parse_single_meeting_docx.
    
    Returns:
        List of dicts with {date, raw_text, sections}
    """
    doc = Document(file_path)
    meetings = []
    current_meeting = None
    current_section = None
    
    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        
        # Check if it's a date line (start of new meeting)
        date_match = re.match(r'^\d{2}\s\w{3}\s\d{2}$', text)
        if date_match:
            # Save previous meeting
            if current_meeting:
                meetings.append(current_meeting)
            # Start new meeting
            current_meeting = {
                'date': parse_date(text),
                'raw_text': text + '\n',
                'sections': {}
            }
            current_section = None
            continue
        
        # Check if it's a section header
        if text in SECTION_HEADERS:
            current_section = text
            if current_meeting:
                current_meeting['sections'][current_section] = []
        
        # Add content to current section
        if current_meeting:
            current_meeting['raw_text'] += text + '\n'
            if current_section:
                current_meeting['sections'][current_section].append(text)
    
    # Don't forget last meeting
    if current_meeting:
        meetings.append(current_meeting)
    
    return meetings


def parse_date(date_str: str):
    """Parse date string like '06 Mar 26' to date object."""
    from datetime import datetime as dt
    return dt.strptime(date_str, '%d %b %y').date()
